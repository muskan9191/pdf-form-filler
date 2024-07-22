from fastapi import FastAPI, Form
from fastapi.responses import FileResponse
from docx import Document
from docx2pdf import convert
import os

app = FastAPI()

def create_filled_docx(input_doc_path, name, address, date, checkbox_items, radio_button_selection, output_path):
    try:
        doc = Document(input_doc_path)
        
        tables_data = {
            1: name,
            3: date,
            5: address,
        }
        cell_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell_count in tables_data:
                        cell.text = tables_data[cell_count]
                    cell_count += 1

        activities = ['Reading', 'Walking', 'Music']
        for i, paragraph in enumerate(doc.paragraphs):
            if 'What are your favourite activities?' in paragraph.text:
                for item in checkbox_items:
                    if item not in activities:
                        for j, p in enumerate(doc.paragraphs[i:]):
                            if '☐Other' in p.text:
                                checkbox_paragraph_index = i + j
                                break
                        doc.paragraphs[checkbox_paragraph_index].text = doc.paragraphs[checkbox_paragraph_index].text.replace('☐Other:', f'☑Other: {item}')
                    else:
                        checkbox_paragraph_index = i + checkbox_items.index(item) + 1
                        doc.paragraphs[checkbox_paragraph_index].text = doc.paragraphs[checkbox_paragraph_index].text.replace('☐', '☑')
            
            if 'What is your favourite activity?' in paragraph.text:
                if radio_button_selection in activities:
                    for j, option in enumerate(activities):
                        if radio_button_selection in option:
                            radio_paragraph_index = i + j + 1
                            doc.paragraphs[radio_paragraph_index].text = doc.paragraphs[radio_paragraph_index].text.replace('◯', '⚫')
                else:
                    for j, p in enumerate(doc.paragraphs[i:]):
                        if '◯Other' in p.text:
                            radio_paragraph_index = i + j
                            doc.paragraphs[radio_paragraph_index].text = doc.paragraphs[radio_paragraph_index].text.replace('◯Other:', f'⚫Other: {radio_button_selection}')
                            break
        
        doc.save(output_path)
    except Exception as e:
        raise Exception(f"An error occurred while filling the data: {str(e)}")


@app.post("/fill-form/")
async def fill_form(
    name: str = Form(...),
    address: str = Form(...),
    date: str = Form(...),
    favourite_activities: str = Form(...),
    favourite_activity: str = Form(...),
):
    try:
        favourite_activities_list = favourite_activities.split(',')
        input_doc_path = "sample_pdf.docx"
        output_doc_path = "filled_form.docx"
        output_pdf_path = "filled_form.pdf"

        create_filled_docx(input_doc_path, name, address, date, favourite_activities_list, favourite_activity, output_doc_path)
        convert(output_doc_path)

        os.remove(output_doc_path)

        return FileResponse(output_pdf_path, media_type='application/pdf', filename=output_pdf_path)
    except Exception as e:
        return {"message": f"An error occurred: {str(e)}"}
